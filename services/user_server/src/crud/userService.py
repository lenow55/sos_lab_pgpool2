from typing import Any
from docx.parts.image import Inches
from docx.shared import RGBColor
from tortoise.contrib.pydantic.base import PydanticModel
from tortoise.query_utils import Prefetch
from src.api.paginated import ListResponse
from src.exceptions.http_exceptions import DuplicateValueException, CustomException, NotFoundException
from tortoise.exceptions import DoesNotExist, IntegrityError
from src.database.models import BaseConfig, User
import src.schemas.user as schemas
import uuid as uuid_pkg

from io import BytesIO
from docx import Document as DocumentConstruct
from docx.document import Document
from docx.oxml.simpletypes import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import logging
from src.core import logger as logger_mod
from src.core.schemas import Status
logger = logging.getLogger(__name__)


class UserService():
    async def create_user(self, user_in_obj: schemas.UserCreateInternal) -> schemas.User:
        try:
            user_obj = await User.create(**user_in_obj.model_dump())
        except IntegrityError:
            raise DuplicateValueException()

        out_user: PydanticModel | schemas.User = await schemas.User.from_tortoise_orm(user_obj)
        if isinstance(out_user, schemas.User):
            return out_user
        else:
            logger.error(
                f"Bad user type {type(out_user)}")
            raise CustomException(
                detail="Bad user instance")

    async def get_user(self, user_id: uuid_pkg.UUID, configs_count=None) -> schemas.UserRead:
        try:
            if isinstance(
                    configs_count, int) and configs_count >= 0:
                db_user: PydanticModel | schemas.UserRead = await schemas.UserRead.from_queryset_single(
                    User.get(uuid=user_id).prefetch_related(
                        Prefetch(
                            "base_config",
                            BaseConfig.filter(author_id=user_id).limit(configs_count)
                        )
                    )
                )
            else:
                db_user: PydanticModel | schemas.UserRead = await schemas.UserRead.from_queryset_single(
                    User.get(uuid=user_id).prefetch_related(
                        Prefetch(
                            "base_config",
                            BaseConfig.filter(author_id=user_id)
                        )
                    )
                )
        except DoesNotExist:
            raise NotFoundException(
                detail=f"User {user_id} not found")

        if isinstance(db_user, schemas.UserRead):
            return db_user
        else:
            logger.error(
                f"Bad user type {type(db_user)}"
            )
            raise CustomException(
                detail="Bad user instance")

    async def find_user(self, **kwargs) -> schemas.User:
        try:
            logger.debug(kwargs)
            db_user: PydanticModel | schemas.User\
                = await schemas.User.from_queryset_single(
                    User.get(**kwargs)
                )
        except DoesNotExist:
            raise NotFoundException(
                detail=f"User {kwargs} not found")

        if isinstance(db_user, schemas.User):
            return db_user
        else:
            logger.error(
                f"Bad user type {type(db_user)}"
            )
            raise CustomException(
                detail="Bad user instance")

    async def get_multi_users(
            self,
            offset: int = 0,
            limit: int = 100
    ) -> ListResponse[schemas.User]:
        users = await schemas.User.from_queryset(
            User.all().offset(offset).limit(limit)
        )
        total_count: int = await User.all().count()
        return ListResponse(
            data=users, total_count=total_count)

    async def delete_user(self, user_id: uuid_pkg.UUID) -> Status:
        try:
            db_user = await schemas.User.from_queryset_single(User.get(uuid=user_id))
        except DoesNotExist:
            raise NotFoundException(
                detail=f"User {user_id} not found")

        if not isinstance(db_user, schemas.User):
            logger.error(
                f"Bad user type {type(db_user)}"
            )
            raise CustomException(
                detail="Bad user instance")

        deleted_count = await User.filter(uuid=user_id).delete()
        if not deleted_count:
            raise NotFoundException(detail={
                "error": f"User[{user_id}] not found"})
        return Status(message=f"Deleted user {user_id}")

    async def update_user(
            self,
            user_id: uuid_pkg.UUID,
            user_update_info: schemas.UserUpdate) -> schemas.User:
        try:
            user: User = await User.get(uuid=user_id)
            logger.debug(
                f"Object to update {user_update_info.model_dump()}")
            if len(user_update_info.model_fields_set) != 0:
                user.update_from_dict(
                    user_update_info.model_dump(
                        exclude_unset=True))
                await user.save(update_fields=user_update_info.model_fields_set)
            logger.debug(f"Object Updated {user}")
            return await schemas.User.from_tortoise_orm(user)
        except DoesNotExist:
            raise NotFoundException(
                detail=f"User {user_id} not found")
        except IntegrityError as e:
            raise DuplicateValueException(
                detail=f"Inconsistent state {e}")

    async def generate_report(
            self,
            user_id: uuid_pkg.UUID) -> bytes:
        # Создаем объект Document
        doc: Any | Document = DocumentConstruct()

        # Устанавливаем шрифт для всего текста в Times New Roman
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        # Устанавливаем отступы для всего документа
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        university_name = doc.add_paragraph(
            'МИНОБРНАУКИ РОССИИ')
        university_name.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        university_name.runs[0].font.size = Pt(14)
        university_name.paragraph_format.space_after = Pt(0)

        # Название учебного заведения
        institution_name = doc.add_paragraph(
            'Федеральное государственное бюджетное образовательное учреждение высшего образования')
        institution_name.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        institution_name.runs[0].font.size = Pt(14)
        institution_name.paragraph_format.space_after = Pt(
            0)
        # Пустые строки
        doc.add_paragraph()

        university_name = doc.add_paragraph(
            'НИЖЕГОРОДСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ им. Р.Е. АЛЕКСЕЕВА')
        university_name.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        university_name.runs[0].font.size = Pt(14)
        university_name.paragraph_format.space_after = Pt(0)

        # Название кафедры
        department_name = doc.add_paragraph(
            '\nКафедра «Вычислительные системы и технологии»')
        department_name.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        department_name.runs[0].font.size = Pt(18)
        department_name.paragraph_format.space_after = Pt(0)

        # Пустые строки
        doc.add_paragraph()

    # Информация о лабораторной работе
        lab_report_title = doc.add_heading(
            'Отчет по лабораторной работе № 3', level=1)
        lab_report_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        lab_report_title.runs[0].font.size = Pt(18)
        lab_report_title.runs[0].font.name = 'Times New Roman'

        # или False
        lab_report_title.runs[0].font.bold = False
        lab_report_title.runs[0].font.color.rgb = RGBColor(
            0, 0, 0)  # цвет текста (черный)
        lab_report_title.paragraph_format.space_after = Pt(
            0)

        doc.add_paragraph()

        lab_report_info = doc.add_paragraph(
            '\nпо дисциплине: «Сервис ориентированные системы»\n\n«Спецификация OpenAPI 3.0»')
        lab_report_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        lab_report_info.runs[0].font.size = Pt(16)
        lab_report_info.paragraph_format.space_after = Pt(0)
        # Пустые строки
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        # Информация о студенте
        doc.add_paragraph(
            'Выполнил:').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        user_info = doc.add_paragraph(
            f'Студент гр. 20-ПО\nНовиков Илья Александрович')
        user_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        user_info.runs[0].font.size = Pt(14)
        user_info.paragraph_format.space_after = Pt(0)

        # Пустые строки
        doc.add_paragraph()

        # Информация о проверяющем
        doc.add_paragraph(
            'Проверил:').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        checker_info = doc.add_paragraph('Сапожников В.О.')
        checker_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        checker_info.runs[0].font.size = Pt(14)
        checker_info.paragraph_format.space_after = Pt(0)
        # Пустые строки
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        # Информация о месте и годе
        location_info = doc.add_paragraph('Нижний Новгород')
        location_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        year_info = doc.add_paragraph('2023 г.')
        year_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        year_info.runs[0].font.size = Pt(14)
        year_info.paragraph_format.space_after = Pt(0)

        # Информация о используемых инструментах
        instrumenst_info = doc.add_paragraph()
        instrumenst_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        instrumenst_info.add_run(
            'Использованные инструменты разработки:')

        items = [
            {"name": "FastAPI", "url": "https://fastapi.tiangolo.com/"},
            {"name": "TortoiseORM", "url": "https://tortoise.github.io/"},
            {"name": "PostgreSQL", "url": "https://www.postgresql.org/"},
            {"name": "Docker", "url": "https://www.docker.com/"}
        ]
        for item in items:
            instrumenst_info.add_run('\n• ' + item["name"])
            hiperlink = instrumenst_info.add_run(
                item
                ["url"])
            hiperlink.font.color.rgb = RGBColor(
                0, 0, 255)  # Задаем цвет текста
        instrumenst_info.paragraph_format.space_after = Pt(
            16)

        # Информация о пользователе
        user_info = doc.add_paragraph()
        user_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        user_info.add_run(
            'Информация о пользователе и двух первых его конфигурациях:')
        user_model_info = await self.get_user(
            user_id=user_id, configs_count=2
        )
        user_serialised_dict = user_model_info.model_dump(
            mode='json')

        for key, value in user_serialised_dict.items():
            user_info.add_run('\n• ' + key).bold = True
            if isinstance(value, list):
                for index, item in enumerate(value):
                    user_info.add_run(
                        '\n\t- ' + str(index)).bold = True
                    if isinstance(item, dict):
                        for key2, value2 in item.items():
                            user_info.add_run(
                                '\n\t\t> ' + key2).bold = True
                            user_info.add_run(
                                ': ' + str(value2))
            else:
                user_info.add_run(': ' + str(value))
        # Сохраняем содержимое документа в байтовый поток
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        return doc_stream.getvalue()


userService: UserService = UserService()
